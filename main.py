from docx import Document
import json
import os, glob
import re

filename = glob.glob('Input\*.docx')

dicts = {}
table_data = []
paragraph_data = []
vids = {}
field_keys = ['2.4. Шифр НД по сварке', '2.5. Группа основного материала', '2.7. Тип сварного шва',
              '2.8. Тип и вид соединения', '2.11. Положение при сварке', 'арматуры железобетонных конструкций',
              '2.15. Положение осей стержней при сварке', '1.2. Дата рождения',
              '1.3. Место работы (сокращенное наименование)', '1.4. Стаж работы по сварке',
              '1.5. Квалификационный разряд (при наличии)', '2.1. Вид аттестации', '2.2. Способ сварки (наплавки)']
out_data = {}
# Члены комиссии
comm_members = ['2347', '2621330', '2014498', '2014498', ]
RD = {'НГДО': {'п.3': ['111540', '111543', '111544', '111545'], 'п.4': [], 'п.13': ['13', '13', '13']},
      'СК': {'п.1': ['1', '1'], 'п.2': ['2', '2'], 'п.3': ['3', '3']},
      'КО': {'п.1': ['1', '1'], 'п.2': ['2', '2'], 'п.3': ['3', '3']},
      'ГО': {'п.1': ['1', '1'], 'п.2': ['2', '2'], 'п.3': ['3', '3']}}
G = {'НГДО': {'3': ['12', '56'], '4': ['101', '202', '303']}, 'СК': {'7': ['4', '5'], '11': ['6', '7']}}
RAD = {}

document = Document(filename[0])
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                table_data.append(para.text)

for para in document.paragraphs:
    paragraph_data.append(para.text)


# Поиск шифра НД, регламентирующих нормы оценки качества
def searchndkontrkach(sp):
    for i in sp:
        if re.findall('[)]:.+', i):
            nd = re.findall('[)]:.+', i)[0][2:]
            return nd


# Поиск города и инн
def searchinncity(parag):
    for i in parag:
        if re.fullmatch('\d{10}[\t]?', i):
            dicts['inn'] = i
    city = re.findall('\sг.[а-яА-ЯёЁ\s]+', parag[4])
    dicts['city'] = city[0]


# Поиск технических устройств опо
def searchvid(sp):
    mets = []
    for i in range(0, len(sp)):
        if sp[i] == 'производственных объектов (ТУ ОПО)':
            i += 1
            while sp[i] != '2.4. Шифр НД по сварке':
                if len(sp[i]) > 1:
                    mets.append(sp[i])
                i += 1
    formatvid(sp=mets)


# Форматирование ту опо в словарь
def formatvid(sp):
    for i in sp:
        osn = re.findall('[А-Я]+', i)[0]
        dop = ','.join(re.findall('п[.\s]\d+', i))
        vids[osn] = dop


# Поиск данных в заявке по ключевым полям
def searchdatafromfield(field_keys, data):
    for i in field_keys:
        index = data.index(i)
        out_data[i] = data[index + 1]


# Поиск и форматирование ФИО
def searchandsplitfio(data):
    index = data.index('1.1. Фамилия, имя, отчество')
    fio = data[index + 1].split()
    out_data['Фамилия'] = fio[0]
    out_data['Имя'] = fio[1]
    out_data['Отчество'] = fio[2]


# Поиск вида свариваемых деталей
def searchvidsvardet(sp):
    index = sp.index('2.6. Вид свариваемых деталей')
    vidsvardet1 = sp[index + 1].split()
    vidsvardet2 = []
    for i in vidsvardet1:
        if re.fullmatch('[;]?[Т][;]?', i):
            vidsvardet2.append('Т1')
        elif re.fullmatch('[;]?[Л][;]?', i):
            vidsvardet2.append('Л1')
        elif re.fullmatch('[;]?[С][;]?', i):
            vidsvardet2.append('С1')
        else:
            vidsvardet2.append(i)
    vidsvardet = ' '.join(vidsvardet2)
    return vidsvardet


# Поиск диапазонов  диаметров и толщин деталей
def searchtolchanddiam(data):
    keys = ['2.9. Диапазон толщин деталей', '2.10. Диапазон диаметров деталей']
    for i in keys:
        index = data.index(i)
        range_str = data[index + 1]
        range = []
        if 'выше' in range_str:
            range.append(range_str)
            range.append('')
        else:
            range.append(re.findall('\d{1,},\d{1,}', range_str)[0])
            range.append(re.findall('\d{1,},\d{1,}', range_str)[1])
        out_data[i] = range


# Поиск диаметров стержней
def searchrangestersh(sp):
    index = sp.index('2.14. Диапазон диаметров стержней')
    rangestersh1 = sp[index + 1]
    rangestersh = []
    if len(rangestersh1) > 2:
        if 'выше' in rangestersh1:
            rangestersh.append(rangestersh1)
            rangestersh.append('')
        else:
            rangestersh.append(re.findall('\d{1,}[,]?\d*', rangestersh1)[0])
            rangestersh.append(re.findall('\d{1,}[,]?\d*', rangestersh1)[1])

    return rangestersh


# Поиск сварочных материалов
def searchsvarmater(sp):
    index = sp.index('2.12. Сварочные материалы')
    svarmater1 = sp[index + 1].split()
    svarmater2 = []
    for i in svarmater1:
        if re.fullmatch('[;]?[Б][;]?', i):
            svarmater2.append('Б1')
        elif re.fullmatch('[;]?[А][;]?', i):
            svarmater2.append('А1')
        elif re.fullmatch('[;]?[Р][;]?', i):
            svarmater2.append('Р1')
        else:
            svarmater2.append(i)
    svarmater = ' '.join(svarmater2)
    return svarmater


def select(method):
    if 'РАД' in method:
        search_blueprint(RAD, group=dicts['group'])
    elif 'РД' in method:
        search_blueprint(RD, group=dicts['group'])


def search_blueprint(dict, group):
    out_blueprint = []
    for vid in group.keys():
        for numbrs in group[vid].split(','):
            out_blueprint += dict[vid][numbrs]
    dicts['blueprints'] = out_blueprint


searchandsplitfio(data=table_data)
searchdatafromfield(field_keys=field_keys, data=table_data)
searchtolchanddiam(data=table_data)
searchvid(sp=table_data)
searchinncity(parag=paragraph_data)

dicts['fam'] = out_data['Фамилия']
dicts['nam'] = out_data['Имя']
dicts['otch'] = out_data['Отчество']
dicts['bdate'] = out_data['1.2. Дата рождения']
dicts['wplace'] = out_data['1.3. Место работы (сокращенное наименование)'][:-6] + 'Газпром Трансгаз Нижний Новгород'
dicts['staj'] = out_data['1.4. Стаж работы по сварке']
dicts['razr'] = out_data['1.5. Квалификационный разряд (при наличии)']
dicts['vid'] = out_data['2.1. Вид аттестации']
dicts['method'] = re.findall('[А-Я]+', out_data['2.2. Способ сварки (наплавки)'])[0]
dicts['group'] = vids
dicts['osnmatgroup'] = out_data['2.5. Группа основного материала']
dicts['vidsravdet'] = searchvidsvardet(table_data)
dicts['typesvarshov'] = out_data['2.7. Тип сварного шва']
dicts['typeandvidconnect'] = out_data['2.8. Тип и вид соединения']
dicts['tolchrange'] = out_data['2.9. Диапазон толщин деталей']
dicts['diamrange'] = out_data['2.10. Диапазон диаметров деталей']
dicts['rangestersh'] = searchrangestersh(table_data)
dicts['svarposit'] = out_data['2.11. Положение при сварке']
dicts['svarmater'] = searchsvarmater(table_data)
dicts['shifrvd'] = out_data['2.4. Шифр НД по сварке']
dicts['typebygost'] = out_data['арматуры железобетонных конструкций']
dicts['sterzhosposit'] = out_data['2.15. Положение осей стержней при сварке']
dicts['ndkontrkach'] = searchndkontrkach(paragraph_data)
dicts['members'] = comm_members
select(dicts['method'])

print(dicts)
with open('Out/data.txt', 'w') as outfile:
    json.dump(dicts, outfile)

# for file in glob.glob("Input\*"):
#     os.remove(file)

#  рабочая версия
# // ==UserScript==
# // @name         Script1
# // @namespace    http://tampermonkey.net/
# // @version      0.1
# // @description  try to take over the world!
# // @author       You
# // @match        https://ac.naks.ru/ac_personal/
# // @icon         https://www.google.com/s2/favicons?sz=64&domain=naks.ru
# // @grant        GM_registerMenuCommand
# // @run-at       context-menu
# // ==/UserScript==
#
# window.addEventListener("DOMContentLoaded", event => {
#     var dd=document.createElement('input');
#     dd.type="file";
#     dd.id = '4';
#     document.getElementById ('navigation').appendChild (dd);
#    var input1=document.createElement('input');
#    input1.type="button";
#    input1.value = 'Записать1';
#    input1.onclick = readFile;
#    document.getElementById ('navigation').appendChild (input1);
#    var input2=document.createElement('input');
#    input2.type="button";
#    input2.value = 'Записать2';
#    input2.onclick = readFile2;
#    document.getElementById ('navigation').appendChild (input2);
#    var input3=document.createElement('input');
#    input3.type="button";
#    input3.value = 'Записать3';
#    input3.onclick = readFile3;
#    document.getElementById ('navigation').appendChild (input3);
#    var input4=document.createElement('input');
#    input4.type="button";
#    input4.value = 'Записать4';
#    input4.onclick = readFile4;
#    document.getElementById ('navigation').appendChild (input4);
#    var input5=document.createElement('input');
#    input5.type="button";
#    input5.value = 'Записать5';
#    input5.onclick = readFile5;
#    document.getElementById ('navigation').appendChild (input5);
# });
#
# function readFile() {
#   var obj = document.getElementById('4');
#   var file = obj.files[0];
#   var reader = new FileReader();
#   reader.onload = function() {
#      var fam = document.getElementsByName('prop[last_name]')[0];
#      var nam = document.getElementsByName('prop[name]')[0];
#      var otch = document.getElementsByName('prop[second_name]')[0];
#      var wplace = document.getElementsByName('prop[company]')[0];
#      var inn = document.getElementsByName('prop[company_inn]')[0];
#      var city = document.getElementsByName('prop[city]')[0];
#      var zay = document.getElementsByName('fiz')[0];
#      var vid = document.getElementsByName('prop[vid_att]')[0];
#      var spos = document.getElementsByName('prop[svarka]')[0];
#      var prof = document.getElementsByName('prop[position]')[0];
#      prof.value = "Сварщик";
#      zay.value = 1;
#      const loaddata = reader.result;
#      var mydata = JSON.parse(loaddata);
#      fam.value = mydata.fam;
#      nam.value = mydata.nam;
#      otch.value = mydata.otch;
#      wplace.value = mydata.wplace;
#       inn.value = mydata.inn;
#       city.value = mydata.city;
#       if (mydata.vid == 'Дополнительная'){
#           vid.value = 96198;}
#       else if (mydata.vid == 'Периодическая'){
#           vid.value = 96195;}
#       else {};
#       if (mydata.method =='РД'){
#           spos.value = 36;}
#       else if (mydata.method == 'Г'){
#           spos.value = 51;}
#       else if (mydata.method == 'РАД'){
#           spos.value = 34;}
#       else if (mydata.method == 'Т'){
#           spos.value = 2094;}
#       else if (mydata.method == 'МП'){
#           spos.value = 65;}
#       else if (mydata.method == 'МПС'){
#           spos.value = 58;}
#       else if (mydata.method == 'НИ'){
#           spos.value = 40;}
#       else if (mydata.method == 'ЗН'){
#           spos.value = 39;}
#       else if (mydata.method == 'АПГ'){
#           spos.value = 63;}
#       else {};
#
#   }
#   reader.readAsText(file)
# }
#
# function readFile2(){
#     var obj = document.getElementById('4');
#     var file = obj.files[0];
#     var reader = new FileReader();
#     reader.onload = function() {
#         const loaddata = reader.result;
#         var mydata = JSON.parse(loaddata);
#         var vklad = document.getElementsByName('prop[dop_att]')[0];
#         var vid = document.getElementsByName('prop[vid_att]')[0];
#         if ( vid.value == 96198){
#             vklad.value = "В1"};
#         var inputs = document.getElementsByTagName('input');
#         var pao = document.getElementsByName('gazprom1')[0];
#         var jj = window.mydata;
#         var list1 = document.getElementById('child_gtu_1').children;
#         var list2 = list1[0].getElementsByTagName('td');
#         var ko = list2[8].children[0];
#         var go = list2[4].children[0];
#         var ngdo = list2[16].children[0];
#         var sk = list2[32].children[0];
#         var keys_v = Object.keys(mydata.group);
#         for (let i in keys_v){
#             if( keys_v[i] == 'ГО'){
#                 go.checked = true;
#                 if (mydata.group['ГО'].indexOf('п.1') !== -1){
#                     var gop1 = document.getElementById('1_elem_tu_340');
#                     gop1.checked=true;}
#                 if (mydata.group['ГО'].indexOf('п.2') !== -1){
#                     var gop2 = document.getElementById('1_elem_tu_3862');
#                     gop2.checked=true;}
#                 if (mydata.group['ГО'].indexOf('п.2п') !== -1){
#                     var gop2p = document.getElementById('1_elem_tu_341');
#                     gop2p.checked=true;}
#                 if (mydata.group['ГО'].indexOf('п.3') !== -1){
#                     var gop3 = document.getElementById('1_elem_tu_342');
#                     gop3.checked=true;}
#                 if (mydata.group['ГО'].indexOf('п.4') !== -1){
#                     var gop4 = document.getElementById('1_elem_tu_343');
#                     gop4.checked=true;}
#                 }
#             else if (keys_v[i] == 'КО'){
#                 ko.checked = true;
#                 if (mydata.group['КО'].indexOf('п.1') !== -1){
#                     var kop1 = document.getElementById('1_elem_tu_335');
#                     kop1.checked=true;}
#                 if (mydata.group['КО'].indexOf('п.2') !== -1){
#                     var ko2 = document.getElementById('1_elem_tu_336');
#                     ko2.checked=true;}
#                 if (mydata.group['КО'].indexOf('п.3') !== -1){
#                     var ko3 = document.getElementById('1_elem_tu_337');
#                     ko3.checked=true;}
#                 if (mydata.group['КО'].indexOf('п.4') !== -1){
#                     var ko4 = document.getElementById('1_elem_tu_338');
#                     ko4.checked=true;}
#             }
#             else if (keys_v[i] == 'СК'){
#                 sk.checked = true;
#                 if (mydata.group['СК'].indexOf('п.1') !== -1){
#                     var sk1 = document.getElementById('1_elem_tu_386');
#                     sk1.checked=true;}
#                 if (mydata.group['СК'].indexOf('п.2') !== -1){
#                     var sk2 = document.getElementById('1_elem_tu_387');
#                     sk2.checked=true;}
#                 if (mydata.group['СК'].indexOf('п.3') !== -1){
#                     var sk3 = document.getElementById('1_elem_tu_388');
#                     sk3.checked=true;}
#                 if (mydata.group['СК'].indexOf('п.4') !== -1){
#                     var sk4 = document.getElementById('1_elem_tu_389');
#                     sk4.checked=true;}
#             }
#             else if (keys_v[i] == 'НГДО'){
#                 ngdo.checked = true;
#                 pao.checked=true;
#                 if (mydata.group['НГДО'].indexOf('п.3') !== -1){
#                     var ngdo3 = document.getElementById('1_elem_tu_349');
#                     ngdo3.checked=true;}
#                 if (mydata.group['НГДО'].indexOf('п.4') !== -1){
#                     var ngdo4 = document.getElementById('1_elem_tu_350');
#                     ngdo4.checked=true;}
#                 if (mydata.group['НГДО'].indexOf('п.10') !== -1){
#                      var ngdo10 = document.getElementById('1_elem_tu_356');
#                     ngdo10.checked=true;}
#                 if (mydata.group['НГДО'].indexOf('п.13') !== -1){
#                      var ngdo13 = document.getElementById('1_elem_tu_359');
#                     ngdo13.checked=true;}
#             }}
# }
#     reader.readAsText(file);
# }
#
# function readFile3() {
#   var obj = document.getElementById('4');
#   var file = obj.files[0];
#   var reader = new FileReader();
#   reader.onload = function() {
#      var bdate = document.getElementById('prop[birthday]');
#      var staj = document.getElementsByName('prop[long]')[0];
#      var razr = document.getElementsByName('prop[okz]')[0];
#      var shifrntd = document.getElementsByName('prop[shifr_nd]')[0];
#      var vid_c = document.getElementById('change_check_2010128');
#       var vid_tt = document.getElementById('change_check_1173');
#       var vid_lt = document.getElementById('change_check_1132');
#       var vid_t = document.getElementById('change_check_1131');
#       var vid_l = document.getElementById('change_check_1130');
#       var vid_cc = document.getElementById('change_check_80021');
#       var vid_lc = document.getElementById('change_check_2010160');
#       var group_m01 = document.getElementById('change_check_93126');
#       var group_m03 = document.getElementById('change_check_93128');
#       var group_m07 = document.getElementById('change_check_93137');
#       var group_m11 = document.getElementById('change_check_93138');
#       var group_m01m11 = document.getElementById('change_check_1416');
#       var group_m01m03 = document.getElementById('change_check_1402');
#       var type_svar_csh = document.getElementsByName('sm_type[]')[0];
#       var type_svar_ush = document.getElementsByName('sm_type[]')[1];
#       var type_soed_bp = document.getElementsByName('sm_connect[]')[0];
#       var type_soed_sp = document.getElementsByName('sm_connect[]')[1];
#       var type_soed_bz = document.getElementsByName('sm_connect[]')[2];
#       var type_soed_zk = document.getElementsByName('sm_connect[]')[3];
#       var range_tolch_det_min = document.getElementsByName('sm_thick[from]')[0];
#       var range_tolch_det_max = document.getElementsByName('sm_thick[to]')[0];
#       var range_diam_det_min = document.getElementsByName('sm_diametr[from]')[0];
#       var range_diam_det_max = document.getElementsByName('sm_diametr[to]')[0];
#       var range_diam_sterz_min = document.getElementsByName('arm_diametr[from]')[0];
#       var range_diam_sterz_max = document.getElementsByName('arm_diametr[to]')[0];
#       var polozh_os_stersh_b = document.getElementsByName('arm_position[]')[0];
#       var polozh_os_stersh_g = document.getElementsByName('arm_position[]')[1];
#       var name_by_gost = document.getElementById('sm_gost_arm');
#       var polozh_pri_svar_n1 = document.getElementsByName('sm_position[]')[0];
#       var polozh_pri_svar_n2 = document.getElementsByName('sm_position[]')[1];
#       var polozh_pri_svar_g = document.getElementsByName('sm_position[]')[2];
#       var polozh_pri_svar_p1 = document.getElementsByName('sm_position[]')[3];
#       var polozh_pri_svar_p2 = document.getElementsByName('sm_position[]')[4];
#       var polozh_pri_svar_v1 = document.getElementsByName('sm_position[]')[5];
#       var polozh_pri_svar_v2 = document.getElementsByName('sm_position[]')[6];
#       var polozh_pri_svar_n45 = document.getElementsByName('sm_position[]')[7];
#       var svar_mat_a = document.getElementById('but_1');
#       var svar_mat_p = document.getElementById('but_2');
#       var svar_mat_pa = document.getElementById('but_3');
#       var svar_mat_pb = document.getElementById('but_4');
#       var svar_mat_pc = document.getElementById('but_5');
#       var svar_mat_b = document.getElementById('but_6');
#       var prisad_mat = document.getElementsByName('prop[sm_prisadka]')[0];
#       var normat_kach = document.getElementsByName('prop[nd_control]')[0];
#      const loaddata = reader.result;
#      var mydata = JSON.parse(loaddata);
#       bdate.value= mydata.bdate;
#       staj.value = mydata.staj;
#       razr.value = mydata.razr;
#       shifrntd.value = mydata.shifrvd;
#       if (mydata.vidsravdet.indexOf('С1') !== -1){
#           vid_c.checked=true;};
#       if(mydata.vidsravdet.indexOf('Т+Т') !== -1){
#           vid_tt.checked=true;};
#       if (mydata.vidsravdet.indexOf('Л+Т') !== -1){
#           vid_lt.checked=true;};
#       if (mydata.vidsravdet.indexOf('Т1') !== -1){
#           vid_t.checked=true;};
#       if (mydata.vidsravdet.indexOf('Л1') !== -1){
#           vid_l.checked=true;};
#       if (mydata.vidsravdet.indexOf('С+С') !== -1){
#           vid_cc.checked=true;};
#       if (mydata.vidsravdet.indexOf('Л+С') !== -1){
#           vid_lc.checked=true;};
#       if (mydata.osnmatgroup.indexOf('М01') !== -1){
#           group_m01.checked=true;};
#       if (mydata.osnmatgroup.indexOf('М03') !== -1){
#           group_m03.checked=true;};
#       if (mydata.osnmatgroup.indexOf('М07') !== -1){
#           group_m07.checked=true;};
#       if (mydata.osnmatgroup.indexOf('М03+М01') !== -1){
#           group_m01m03.checked=true;};
#       if (mydata.osnmatgroup.indexOf('М01+М11') !== -1){
#           group_m01m11.checked=true;};
#       if (mydata.osnmatgroup.indexOf('М11') !== -1){
#           group_m11.checked=true;};
#       if (mydata.typesvarshov.indexOf('СШ') !== -1){
#           type_svar_csh.checked=true;};
#       if (mydata.typesvarshov.indexOf('УШ') !== -1){
#           type_svar_ush.checked=true;};
#       if (mydata.typeandvidconnect.indexOf('бп') !== -1){
#           type_soed_bp.checked=true;};
#       if (mydata.typeandvidconnect.indexOf('сп') !== -1){
#           type_soed_sp.checked=true;};
#       if (mydata.typeandvidconnect.indexOf('бз') !== -1){
#           type_soed_bz.checked=true;};
#       if (mydata.typeandvidconnect.indexOf('зк') !== -1){
#           type_soed_zk.checked=true;};
#        if (mydata.tolchrange.length > 0){
#            range_tolch_det_min.value = mydata.tolchrange[0];
#            range_tolch_det_max.value = mydata.tolchrange[1];
#                                         };
#       if (mydata.diamrange.length > 0){
#           range_diam_det_min.value = mydata.diamrange[0];
#           range_diam_det_max.value = mydata.diamrange[1];
#                                        };
#       if (mydata.rangestersh.length > 0){
#           range_diam_sterz_min.value = mydata.rangestersh[0];
#           range_diam_sterz_max.value = mydata.rangestersh[1];
#                                         };
#       if (mydata.svarposit.indexOf('Н1') !== -1){
#           polozh_pri_svar_n1.checked=true;};
#       if (mydata.svarposit.indexOf('Н2') !== -1){
#           polozh_pri_svar_n2.checked=true;};
#       if (mydata.svarposit.indexOf('Г') !== -1){
#           polozh_pri_svar_g.checked=true;};
#       if (mydata.svarposit.indexOf('П1') !== -1){
#           polozh_pri_svar_p1.checked=true;};
#       if (mydata.svarposit.indexOf('П2') !== -1){
#           polozh_pri_svar_p2.checked=true;};
#       if (mydata.svarposit.indexOf('В1') !== -1){
#           polozh_pri_svar_v1.checked=true;};
#       if (mydata.svarposit.indexOf('В2') !== -1){
#           polozh_pri_svar_v2.checked=true;};
#       if (mydata.svarposit.indexOf('Н45') !== -1){
#           polozh_pri_svar_n45.checked=true;};
#       if (mydata.method != 'РАД'){
#           if (mydata.svarmater.indexOf('А1') !== -1){
#               svar_mat_a.click();};
#           if (mydata.svarmater.indexOf('Р1') !== -1){
#               svar_mat_p.click();};
#           if (mydata.svarmater.indexOf('РА') !== -1){
#               svar_mat_pa.click();};
#           if (mydata.svarmater.indexOf('РБ') !== -1){
#               svar_mat_pb.click();};
#           if (mydata.svarmater.indexOf('РЦ') !== -1){
#               svar_mat_pc.click();};
#           if (mydata.svarmater.indexOf('Б1') !== -1){
#               svar_mat_b.click();};}
#           else {prisad_mat.value = mydata.svarmater };
#       name_by_gost.value = mydata.typebygost;
#       if (mydata.sterzhosposit.indexOf('В') !== -1){
#           polozh_os_stersh_b.checked=true;};
#       if (mydata.sterzhosposit.indexOf('Г') !== -1){
#           polozh_os_stersh_g.checked=true;};
#       normat_kach.value = mydata.ndkontrkach;
#         };
#   reader.readAsText(file)
# }
#
# function readFile4() {
#   var obj = document.getElementById('4');
#   var file = obj.files[0];
#   var reader = new FileReader();
#   reader.onload = function() {
#      var lst_members = document.getElementsByTagName('input');
#      const loaddata = reader.result;
#      var mydata = JSON.parse(loaddata);
#       var memb1 = mydata.members[0];
#       var memb2 = mydata.members[1];
#       var memb3 = mydata.members[2];
#       var preds = mydata.members[3];
#       for (let i in lst_members){
#       if (lst_members[i].type == 'checkbox' && lst_members[i].value == memb1 ){
#           lst_members[i].checked = true}
#       else if (lst_members[i].type == 'checkbox' && lst_members[i].value == memb2 ){
#           lst_members[i].checked = true}
#       else if (lst_members[i].type == 'checkbox' && lst_members[i].value == memb3 ){
#           lst_members[i].checked = true}
#       else if (lst_members[i].type == 'radio' && lst_members[i].value == preds ){
#           lst_members[i].checked = true}
#       };
#   }
#   reader.readAsText(file)
# }
#
# function readFile5() {
#   var obj = document.getElementById('4');
#   var file = obj.files[0];
#   var reader = new FileReader();
#   reader.onload = function() {
#      var list_data = ['112272', '91723', '82277']
#      const loaddata = reader.result;
#      var mydata = JSON.parse(loaddata);
#       (function myLoop (i) {
#           setTimeout(function () {var tbl = document.getElementById('child_shablon');
#                      var but_rec = tbl.getElementsByTagName('input')[0];
#                      var select_bp = document.getElementsByName('from_kss_shablon')[0];
#                      console.log(i);
#                      select_bp.value = list_data[i];
#                      but_rec.click();
#                      i--;
#                      if (i>=0) myLoop(i);
#                      }, 1500)
#       })(list_data.length-1);
#
#       }
#   reader.readAsText(file);
# }
